"""
Document parsing.

Extracts text while retaining enough metadata (page numbers for PDFs,
heading/section context for DOCX) to support citations later in the
pipeline. Each parser returns a list of ExtractedBlock, which the chunker
then groups into chunks.
"""
from dataclasses import dataclass

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.errors import DocumentProcessingFailedError


@dataclass
class ExtractedBlock:
    text: str
    page_number: int | None = None
    section: str | None = None


def parse_pdf(file_path: str) -> list[ExtractedBlock]:
    blocks: list[ExtractedBlock] = []
    try:
        with fitz.open(file_path) as doc:
            for page_index, page in enumerate(doc, start=1):
                text = page.get_text("text").strip()
                if text:
                    blocks.append(ExtractedBlock(text=text, page_number=page_index))
    except Exception as exc:
        raise DocumentProcessingFailedError(f"Could not read PDF: {exc}") from exc

    if not blocks:
        raise DocumentProcessingFailedError("No extractable text found in this PDF (it may be scanned/image-only).")
    return blocks


def parse_docx(file_path: str) -> list[ExtractedBlock]:
    blocks: list[ExtractedBlock] = []
    try:
        doc = DocxDocument(file_path)
        current_section = None
        buffer: list[str] = []

        def flush():
            if buffer:
                blocks.append(ExtractedBlock(text="\n".join(buffer).strip(), section=current_section))
                buffer.clear()

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading") or style == "title":
                flush()
                current_section = text
                buffer.append(text)
            else:
                buffer.append(text)
        flush()

        # Tables: extract as their own blocks so clause/limit tables aren't lost
        for table_index, table in enumerate(doc.tables):
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows_text.append(" | ".join(cells))
            if rows_text:
                blocks.append(
                    ExtractedBlock(
                        text="\n".join(rows_text),
                        section=current_section or f"Table {table_index + 1}",
                    )
                )
    except DocumentProcessingFailedError:
        raise
    except Exception as exc:
        raise DocumentProcessingFailedError(f"Could not read DOCX: {exc}") from exc

    if not blocks:
        raise DocumentProcessingFailedError("No extractable text found in this document.")
    return blocks


def parse_txt(file_path: str) -> list[ExtractedBlock]:
    try:
        with open(file_path, encoding="utf-8", errors="replace") as f:
            text = f.read().strip()
    except Exception as exc:
        raise DocumentProcessingFailedError(f"Could not read text file: {exc}") from exc

    if not text:
        raise DocumentProcessingFailedError("The text file is empty.")
    return [ExtractedBlock(text=text)]


def parse_document(file_path: str, mime_type: str) -> list[ExtractedBlock]:
    if mime_type == "application/pdf":
        return parse_pdf(file_path)
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return parse_docx(file_path)
    if mime_type == "text/plain":
        return parse_txt(file_path)
    raise DocumentProcessingFailedError(f"Unsupported mime type: {mime_type}")
